import zipfile
from pathlib import Path

from . import register
from .models import DocumentImage, IMAGE_MIME, NO_TEXT_HINT, ParsedCase

try:
    from docx import Document
except ImportError:
    Document = None


@register(".docx")
def parse_docx(file_path: str) -> list[ParsedCase]:
    if Document is None:
        raise ImportError("解析 docx 需要安装 python-docx: pip install python-docx")

    doc = Document(file_path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)

    images = []
    with zipfile.ZipFile(file_path) as zf:
        for name in zf.namelist():
            if not name.startswith("word/media/") or name.endswith("/"):
                continue
            mime = IMAGE_MIME.get(Path(name).suffix.lstrip(".").lower())
            if mime:
                images.append(DocumentImage(zf.read(name), mime, Path(name).name))

    if not text.strip() and not images:
        return []
    if not text.strip():
        text = NO_TEXT_HINT
    return [ParsedCase(text, images)]
